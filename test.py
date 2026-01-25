from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Typography Showcase', 0)

# English Section
doc.add_heading('English Styles', level=1)
fonts = ["Times New Roman", "Arial", "Courier New", "Comic Sans MS"]
for font in fonts:
    p = doc.add_paragraph(f"This is a sample of {font} font.")
    p.runs[0].font.name = font
    p.runs[0].font.size = Pt(14)

# Marathi Section
doc.add_heading('Marathi Styles', level=1)
marathi_text = "मराठी भाषेचे सौंदर्य आणि विविध फॉन्ट शैली."
# Note: These fonts must be installed on your system to show up correctly
m_fonts = ["Latha", "Nirmala UI", "Mangal"]
for m_font in m_fonts:
    p = doc.add_paragraph(f"{m_font}: {marathi_text}")
    p.runs[0].font.name = m_font
    p.runs[0].font.size = Pt(14)

doc.save('FontShowcase.docx')
